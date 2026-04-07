"""
Text Processing Utilities.

Handles document text extraction (PDF/DOCX), text cleaning/normalization,
and intelligent, paragraph-aware chunking with token-based sizing.
"""

import logging
import re
from typing import List

import fitz  # PyMuPDF
import tiktoken
from docx import Document as DocxDocument

from app.core.config import settings

logger = logging.getLogger(__name__)


class TextProcessor:
    """Utility class for text extraction, cleaning, and chunking."""

    # ──────────────────── Text Extraction ────────────────────

    @staticmethod
    def extract_text(file_path: str, file_type: str) -> str:
        """
        Extract text from a PDF or DOCX file.

        Args:
            file_path: Path to the file on disk.
            file_type: File extension ('pdf' or 'docx').

        Returns:
            Extracted raw text.

        Raises:
            ValueError: If file type is unsupported.
        """
        file_type = file_type.lower().strip(".")

        if file_type == "pdf":
            return TextProcessor._extract_pdf(file_path)
        elif file_type in ("docx", "doc"):
            return TextProcessor._extract_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

    @staticmethod
    def _extract_pdf(file_path: str) -> str:
        """Extract text from a PDF file using PyMuPDF."""
        text_parts: List[str] = []
        try:
            doc = fitz.open(file_path)
            for page_num, page in enumerate(doc):
                page_text = page.get_text("text")
                if page_text.strip():
                    text_parts.append(page_text)
            doc.close()
            logger.info(f"Extracted text from {len(text_parts)} PDF pages")
        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}")
            raise
        return "\n\n".join(text_parts)

    @staticmethod
    def _extract_docx(file_path: str) -> str:
        """Extract text from a DOCX file using python-docx."""
        text_parts: List[str] = []
        try:
            doc = DocxDocument(file_path)
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        text_parts.append(row_text)

            logger.info(f"Extracted {len(text_parts)} paragraphs from DOCX")
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {e}")
            raise
        return "\n\n".join(text_parts)

    # ──────────────────── Text Cleaning ────────────────────

    @staticmethod
    def clean_text(text: str) -> str:
        """
        Clean and normalize extracted text.

        Removes excessive whitespace, special characters, and
        normalizes line breaks while preserving paragraph structure.

        Args:
            text: Raw extracted text.

        Returns:
            Cleaned text.
        """
        # Replace tabs with spaces
        text = text.replace("\t", " ")

        # Normalize line endings
        text = text.replace("\r\n", "\n").replace("\r", "\n")

        # Remove page numbers and headers/footers patterns
        text = re.sub(r"^\s*Page\s+\d+\s*(of\s+\d+)?\s*$", "", text, flags=re.MULTILINE)

        # Remove excessive whitespace within lines
        text = re.sub(r"[ ]{3,}", " ", text)

        # Normalize paragraph breaks (3+ newlines → 2)
        text = re.sub(r"\n{3,}", "\n\n", text)

        # Remove leading/trailing whitespace from each line
        lines = [line.strip() for line in text.split("\n")]
        text = "\n".join(lines)

        # Remove leading/trailing whitespace from entire text
        text = text.strip()

        return text

    # ──────────────────── Token Counting ────────────────────

    @staticmethod
    def count_tokens(text: str) -> int:
        """Count the number of tokens in a text string using tiktoken."""
        try:
            encoding = tiktoken.get_encoding("cl100k_base")
            return len(encoding.encode(text))
        except Exception:
            # Fallback: rough estimate (1 token ≈ 4 chars)
            return len(text) // 4

    # ──────────────────── Smart Chunking ────────────────────

    @staticmethod
    def chunk_text(
        text: str,
        chunk_size: int = settings.CHUNK_SIZE,
        chunk_overlap: int = settings.CHUNK_OVERLAP,
    ) -> List[str]:
        """
        Split text into overlapping chunks with semantic awareness.

        Strategy:
        1. Split text into paragraphs (natural semantic boundaries).
        2. Accumulate paragraphs into chunks until the token limit is reached.
        3. Apply overlap by including the last few sentences from the
           previous chunk in the next chunk.

        Args:
            text: Cleaned text to chunk.
            chunk_size: Target chunk size in tokens (500-800).
            chunk_overlap: Overlap size in tokens (100-150).

        Returns:
            List of text chunks.
        """
        # Split into paragraphs (semantic boundaries)
        paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]

        if not paragraphs:
            # Fallback: split by single newlines
            paragraphs = [p.strip() for p in text.split("\n") if p.strip()]

        if not paragraphs:
            logger.warning("No paragraphs found, returning text as single chunk")
            return [text] if text.strip() else []

        chunks: List[str] = []
        current_chunk_parts: List[str] = []
        current_token_count = 0

        for paragraph in paragraphs:
            para_tokens = TextProcessor.count_tokens(paragraph)

            # If a single paragraph exceeds the chunk size, split it further
            if para_tokens > chunk_size:
                # Flush current chunk if non-empty
                if current_chunk_parts:
                    chunks.append("\n\n".join(current_chunk_parts))
                    current_chunk_parts = []
                    current_token_count = 0

                # Split large paragraph by sentences
                sub_chunks = TextProcessor._split_large_paragraph(
                    paragraph, chunk_size, chunk_overlap
                )
                chunks.extend(sub_chunks)
                continue

            # Check if adding this paragraph exceeds the limit
            if current_token_count + para_tokens > chunk_size:
                # Save current chunk
                if current_chunk_parts:
                    chunk_text = "\n\n".join(current_chunk_parts)
                    chunks.append(chunk_text)

                    # Create overlap: take the last part of the current chunk
                    overlap_text = TextProcessor._get_overlap(
                        current_chunk_parts, chunk_overlap
                    )
                    current_chunk_parts = [overlap_text] if overlap_text else []
                    current_token_count = TextProcessor.count_tokens(overlap_text) if overlap_text else 0

            current_chunk_parts.append(paragraph)
            current_token_count += para_tokens

        # Don't forget the last chunk
        if current_chunk_parts:
            chunks.append("\n\n".join(current_chunk_parts))

        logger.info(
            f"Chunked text into {len(chunks)} chunks "
            f"(target size: {chunk_size} tokens, overlap: {chunk_overlap} tokens)"
        )

        return chunks

    @staticmethod
    def _split_large_paragraph(
        paragraph: str,
        chunk_size: int,
        chunk_overlap: int,
    ) -> List[str]:
        """
        Split a very large paragraph into smaller chunks by sentences.

        Args:
            paragraph: Large paragraph text.
            chunk_size: Target chunk size in tokens.
            chunk_overlap: Overlap in tokens.

        Returns:
            List of sub-chunks.
        """
        # Split by sentence boundaries
        sentences = re.split(r"(?<=[.!?])\s+", paragraph)
        chunks: List[str] = []
        current_sentences: List[str] = []
        current_tokens = 0

        for sentence in sentences:
            sentence_tokens = TextProcessor.count_tokens(sentence)

            if current_tokens + sentence_tokens > chunk_size and current_sentences:
                chunks.append(" ".join(current_sentences))

                # Calculate overlap
                overlap_sentences: List[str] = []
                overlap_tokens = 0
                for s in reversed(current_sentences):
                    s_tokens = TextProcessor.count_tokens(s)
                    if overlap_tokens + s_tokens > chunk_overlap:
                        break
                    overlap_sentences.insert(0, s)
                    overlap_tokens += s_tokens

                current_sentences = overlap_sentences
                current_tokens = overlap_tokens

            current_sentences.append(sentence)
            current_tokens += sentence_tokens

        if current_sentences:
            chunks.append(" ".join(current_sentences))

        return chunks

    @staticmethod
    def _get_overlap(parts: List[str], overlap_tokens: int) -> str:
        """
        Get the last N tokens worth of text from chunk parts for overlap.

        Args:
            parts: List of paragraph parts from the current chunk.
            overlap_tokens: Desired overlap in tokens.

        Returns:
            Overlap text string.
        """
        overlap_parts: List[str] = []
        tokens_so_far = 0

        for part in reversed(parts):
            part_tokens = TextProcessor.count_tokens(part)
            if tokens_so_far + part_tokens > overlap_tokens:
                # Take partial text from this part by sentences
                sentences = re.split(r"(?<=[.!?])\s+", part)
                for sent in reversed(sentences):
                    sent_tokens = TextProcessor.count_tokens(sent)
                    if tokens_so_far + sent_tokens > overlap_tokens:
                        break
                    overlap_parts.insert(0, sent)
                    tokens_so_far += sent_tokens
                break
            overlap_parts.insert(0, part)
            tokens_so_far += part_tokens

        return " ".join(overlap_parts)
